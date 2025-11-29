import os
import concurrent.futures
from typing import List, Dict, Generator

# File Processing
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
try:
    import docx
except ImportError:
    docx = None

# LangChain Imports
from langchain_core.documents import Document
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings

# Sentence Transformers Fallback
try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    SentenceTransformer = None

# ==========================================
# 1. TEXT EXTRACTION
# ==========================================

def extract_text_single_file(file_path: str) -> List[Dict]:
    """Extracts text from a single file. Used by the thread pool."""
    pages = []
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf" and pdfplumber:
            with pdfplumber.open(file_path) as pdf:
                for i, p in enumerate(pdf.pages):
                    text = p.extract_text() or ""
                    if text.strip():
                        pages.append({"page": i + 1, "text": text})
        
        elif ext in [".docx", ".doc"] and docx:
            doc = docx.Document(file_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            if full_text.strip():
                pages.append({"page": 1, "text": full_text})
        
        elif ext in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                pages.append({"page": 1, "text": f.read()})
                
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
    
    return pages

# ==========================================
# 2. EMBEDDING STRATEGY
# ==========================================

class LocalSentenceEmbedding:
    def __init__(self, model="all-MiniLM-L6-v2"):
        if SentenceTransformer is None:
            raise ImportError("sentence-transformers not installed.")
        self.m = SentenceTransformer(model)

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return [v.tolist() for v in self.m.encode(texts, convert_to_numpy=True)]

    def embed_query(self, text: str) -> List[float]:
        return self.m.encode([text], convert_to_numpy=True)[0].tolist()

def get_embeddings_function():
    api_key = os.environ.get("GOOGLE_API_KEY")
    if api_key:
        try:
            return GoogleGenerativeAIEmbeddings(
                model="models/embedding-001",
                google_api_key=api_key
            )
        except Exception:
            pass
    return LocalSentenceEmbedding()

def get_llm(model_name="gemini-2.5-flash"):
    """
    Creates LLM with specific fallback logic for Gemini 2.x models.
    """
    api_key = os.environ.get("GOOGLE_API_KEY")
    
    # Try the requested model first
    try:
        return ChatGoogleGenerativeAI(
            model=model_name,
            google_api_key=api_key,
            temperature=0.3,
            convert_system_message_to_human=True
        )
    except Exception:
        pass
    
    # Fallback 1: Gemini 2.0 Flash
    try:
        print(f"Warning: Could not load {model_name}, trying gemini-2.0-flash")
        return ChatGoogleGenerativeAI(
            model="gemini-2.0-flash",
            google_api_key=api_key,
            temperature=0.3,
            convert_system_message_to_human=True
        )
    except Exception:
        pass

    # Fallback 2: Generic latest alias
    print("Warning: Specific models failed, trying gemini-flash-latest")
    return ChatGoogleGenerativeAI(
        model="gemini-flash-latest",
        google_api_key=api_key,
        temperature=0.3,
        convert_system_message_to_human=True
    )

# ==========================================
# 3. INGESTION LOGIC
# ==========================================

def process_and_index(file_paths: List[str], persist_directory="chroma_db"):
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    all_docs = []

    with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
        future_to_file = {executor.submit(extract_text_single_file, fp): fp for fp in file_paths}
        
        for future in concurrent.futures.as_completed(future_to_file):
            fp = future_to_file[future]
            try:
                pages = future.result()
                for p in pages:
                    chunks = text_splitter.split_text(p['text'])
                    for i, chunk in enumerate(chunks):
                        all_docs.append(Document(
                            page_content=chunk,
                            metadata={
                                "source": os.path.basename(fp),
                                "page": p['page'],
                            }
                        ))
            except Exception as e:
                print(f"Failed to process {fp}: {e}")

    if not all_docs:
        return {"success": False, "message": "No text extracted from files."}

    emb_fn = get_embeddings_function()
    try:
        vectordb = Chroma.from_documents(
            documents=all_docs,
            embedding=emb_fn,
            persist_directory=persist_directory
        )
        vectordb.persist()
        return {"success": True, "message": f"Successfully indexed {len(all_docs)} chunks."}
    except Exception as e:
        return {"success": False, "message": str(e)}

# ==========================================
# 4. INTELLIGENT AGENT LOGIC
# ==========================================

def rewrite_query(original_query: str, chat_history: List, model_name: str) -> str:
    """
    Feature 2: Uses a cheap LLM call to rewrite vague queries into specific ones.
    """
    llm = get_llm(model_name)
    
    history_context = ""
    if chat_history:
        # Take last 2 exchanges to keep context relevant but concise
        last_exchange = chat_history[-2:] 
        for role, text in last_exchange:
            history_context += f"{role}: {text}\n"

    prompt = (
        "You are an AI query refiner. Rewrite the user's query to be optimal for vector retrieval.\n"
        "1. Remove conversational filler.\n"
        "2. Add specific keywords from context if the query is vague (e.g., 'what about cost?').\n"
        "3. Output ONLY the rewritten query text.\n\n"
        f"Chat Context:\n{history_context}\n"
        f"Original Query: {original_query}\n"
        "Refined Query:"
    )
    
    try:
        response = llm.invoke(prompt)
        return response.content.strip()
    except Exception:
        return original_query 

def retrieve_documents(query: str, persist_directory="chroma_db") -> List[Document]:
    """
    Feature 1 (Part A): Retrieves raw documents for the reasoning trace.
    """
    emb_fn = get_embeddings_function()
    try:
        vectordb = Chroma(persist_directory=persist_directory, embedding_function=emb_fn)
        retriever = vectordb.as_retriever(search_kwargs={"k": 5})
        try:
            return retriever.invoke(query)
        except AttributeError:
            return retriever.get_relevant_documents(query)
    except Exception as e:
        print(f"Retrieval Error: {e}")
        return []

def stream_answer(
    question: str, 
    context_docs: List[Document], 
    chat_history: List, 
    model_name: str,
    complexity_level: str = "Standard"
) -> Generator:
    """
    Feature 6: Generates answer with specific complexity level.
    """
    context_text = "\n\n".join([
        f"Source: {d.metadata.get('source', 'Unknown')} (Pg {d.metadata.get('page', '?')})\nContent: {d.page_content}" 
        for d in context_docs
    ])
    
    if not context_text:
        yield "I couldn't find any relevant information in the uploaded documents."
        return

    history_text = "\n".join([f"{role}: {text}" for role, text in chat_history[-4:]])
    
    audience_instruction = "Provide a standard, balanced explanation."
    if complexity_level == "Novice":
        audience_instruction = "Explain this like I am 5 years old. Use simple analogies and avoid jargon."
    elif complexity_level == "Expert":
        audience_instruction = "Provide a highly technical, detailed analysis. Assume the user is a subject matter expert."

    system_prompt = (
        "You are a Knowledge Base Agent. Use the Context below to answer the user's question.\n"
        f"Constraint: {audience_instruction}\n"
        "If the answer isn't in the context, admit it.\n\n"
        f"CONTEXT:\n{context_text}\n\n"
        f"CHAT HISTORY:\n{history_text}\n\n"
        f"USER QUESTION: {question}\n\n"
        "ANSWER:"
    )

    llm = get_llm(model_name)
    try:
        for chunk in llm.stream(system_prompt):
            if chunk.content:
                yield chunk.content
    except Exception as e:
        yield f"Error generating response: {e}"

# ==========================================
# 5. KNOWLEDGE GRAPH GENERATOR (NEW)
# ==========================================
def generate_knowledge_graph(context_docs: List[Document], model_name: str) -> str:
    """
    Feature 4: Extracts entities and relationships for Graphviz.
    Returns: String in DOT format.
    """
    # Use a limited sample of text to avoid token limits and keep the graph focused
    text_sample = "\n".join([d.page_content[:500] for d in context_docs[:3]]) 
    llm = get_llm(model_name)
    
    prompt = (
        "You are a Knowledge Graph extraction engine. Analyze the text below and extract core entities and relationships.\n"
        "Output strictly in DOT format for Graphviz. Do not include ```graphviz``` or ```dot``` markers. Just the code.\n"
        "Format: digraph G { rankdir=LR; NodeA -> NodeB [label=\"relationship\"]; ... }\n"
        "Keep node names short (max 3 words). Limit to the top 10 most important relationships.\n\n"
        f"Text to Analyze:\n{text_sample}"
    )
    
    try:
        response = llm.invoke(prompt)
        # Clean up code blocks if the LLM includes them
        dot_code = response.content.replace("```dot", "").replace("```graphviz", "").replace("```", "").strip()
        
        if "digraph" not in dot_code:
            return 'digraph G { "No Graph" -> "Generated"; }'
        return dot_code
    except Exception as e:
        return f'digraph G {{ "Error" -> "{str(e)}"; }}'
