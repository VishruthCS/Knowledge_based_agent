# ingest.py
"""
Robust ingestion + indexing + querying with:
- Gemini embeddings (if quota available)
- Automatic fallback to local sentence-transformers embeddings
- No dependency on langchain.chains
- Compatible with langchain-core>=0.3
"""

import os
import traceback
from typing import List, Dict, Any

# OCR + file reading
import pytesseract
from PIL import Image
try:
    import pdfplumber
except Exception:
    pdfplumber = None
try:
    from pdf2image import convert_from_path
except Exception:
    convert_from_path = None
try:
    import docx
except Exception:
    docx = None

# LangChain modern imports
try:
    from langchain_core.documents import Document
except Exception:
    from langchain.schema import Document  # fallback for very old LC installs

from langchain_text_splitters import CharacterTextSplitter
from langchain_community.vectorstores import Chroma

# Gemini LLM + embeddings
from langchain_google_genai import (
    ChatGoogleGenerativeAI,
    GoogleGenerativeAIEmbeddings,
)

# Sentence-transformers fallback
try:
    from sentence_transformers import SentenceTransformer
except Exception:
    SentenceTransformer = None


# ---------------------------------------------------------------------
# TEXT EXTRACTION HELPERS
# ---------------------------------------------------------------------
def extract_text_from_pdf(pdf_path, poppler_path=None):
    pages = []
    if pdfplumber is None:
        return [{"page": 1, "text": ""}]

    try:
        with pdfplumber.open(pdf_path) as pdf:
            for i, p in enumerate(pdf.pages):
                t = p.extract_text() or ""
                pages.append({"page": i + 1, "text": t})
    except Exception:
        pages = []

    # If text too small, fallback to OCR
    total = sum(len(p["text"]) for p in pages)
    if total < 50 and convert_from_path is not None:
        try:
            imgs = convert_from_path(pdf_path, dpi=200, poppler_path=poppler_path)
            out = []
            for i, img in enumerate(imgs):
                txt = pytesseract.image_to_string(img)
                out.append({"page": i + 1, "text": txt})
            return out
        except Exception:
            pass

    return pages


def extract_text_from_docx(path):
    if docx is None:
        return [{"page": 1, "text": ""}]
    try:
        d = docx.Document(path)
        full = "\n".join([p.text for p in d.paragraphs])
        return [{"page": 1, "text": full}]
    except Exception:
        return [{"page": 1, "text": ""}]


def extract_text_from_txt(path):
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return [{"page": 1, "text": f.read()}]
    except Exception:
        return [{"page": 1, "text": ""}]


def extract_text(path, poppler_path=None):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path, poppler_path)
    if ext in [".docx", ".doc"]:
        return extract_text_from_docx(path)
    if ext in [".txt", ".md"]:
        return extract_text_from_txt(path)

    # image fallback
    try:
        img = Image.open(path)
        txt = pytesseract.image_to_string(img)
        return [{"page": 1, "text": txt}]
    except Exception:
        return [{"page": 1, "text": ""}]


# ---------------------------------------------------------------------
# CHUNKING
# ---------------------------------------------------------------------
def create_documents_from_files(files, chunk_size=800, chunk_overlap=150, poppler_path=None):
    splitter = CharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    out = []

    for fp in files:
        pages = extract_text(fp, poppler_path)
        for p in pages:
            txt = (p["text"] or "").strip()
            if not txt:
                continue
            chunks = splitter.split_text(txt)
            for i, ch in enumerate(chunks):
                out.append(
                    Document(
                        page_content=ch,
                        metadata={
                            "source": os.path.basename(fp),
                            "file_path": fp,
                            "page": p["page"],
                            "chunk_id": f"{os.path.basename(fp)}_p{p['page']}_c{i}",
                        },
                    )
                )
    return out


# ---------------------------------------------------------------------
# EMBEDDINGS — AUTO FALLBACK
# ---------------------------------------------------------------------
class LocalSentenceEmbedding:
    """Wrapper that matches embed_documents / embed_query for Chroma."""
    def __init__(self, model="all-MiniLM-L6-v2"):
        if SentenceTransformer is None:
            raise ImportError("sentence-transformers not installed.")
        self.m = SentenceTransformer(model)

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        arr = self.m.encode(texts, convert_to_numpy=True, show_progress_bar=False)
        return [v.tolist() for v in arr]

    def embed_query(self, text: str) -> List[float]:
        arr = self.m.encode([text], convert_to_numpy=True, show_progress_bar=False)
        return arr[0].tolist()


def get_embeddings():
    """Try Gemini embeddings → fallback to local sentence-transformers."""
    key = os.environ.get("GOOGLE_API_KEY")

    if key:
        try:
            emb = GoogleGenerativeAIEmbeddings(
                model="models/embedding-001",
                google_api_key=key
            )
            # Smoke test
            emb.embed_documents(["test"])
            print("Using Gemini embeddings.")
            return emb
        except Exception as e:
            print("Gemini embeddings failed, falling back to local:", e)
            traceback.print_exc()

    if SentenceTransformer is None:
        raise RuntimeError("No Gemini quota and no sentence-transformers installed.")

    print("Using local MiniLM embeddings.")
    return LocalSentenceEmbedding()


# ---------------------------------------------------------------------
# INDEXING
# ---------------------------------------------------------------------
def index_documents(docs, persist_directory="chroma_db"):
    if not docs:
        return {"success": False, "message": "No documents found."}

    emb = get_embeddings()

    try:
        vectordb = Chroma.from_documents(docs, emb, persist_directory=persist_directory)
        vectordb.persist()
        return {"success": True, "message": f"Indexed {len(docs)} chunks."}
    except Exception as e:
        return {"success": False, "message": f"Indexing failed: {e}"}


def process_and_index(files, persist_directory="chroma_db", poppler_path=None):
    docs = create_documents_from_files(files, poppler_path=poppler_path)
    if not docs:
        return {"success": False, "message": "No extracted text."}
    return index_documents(docs, persist_directory)


# ---------------------------------------------------------------------
# QUERY
# ---------------------------------------------------------------------
def query_index(question, persist_directory="chroma_db", top_k=5):
    key = os.environ.get("GOOGLE_API_KEY")
    emb = get_embeddings()

    try:
        vectordb = Chroma(persist_directory=persist_directory, embedding_function=emb)
    except Exception as e:
        return {"answer": "", "sources": [], "error": f"Cannot load index: {e}"}

    # Retrieve
    retriever = vectordb.as_retriever(search_kwargs={"k": top_k})
    try:
        docs = retriever.get_relevant_documents(question)
    except Exception:
        docs = vectordb.similarity_search(question, k=top_k)

    # Build context
    snippets = []
    sources = []

    for d in docs:
        meta = d.metadata or {}
        txt = d.page_content[:1000]
        snippets.append(
            f"Source: {meta.get('source')} (page {meta.get('page')})\n{txt}"
        )
        sources.append({
            "filename": meta.get("source"),
            "page": meta.get("page"),
            "file_path": meta.get("file_path"),
            "chunk_id": meta.get("chunk_id"),
            "text": txt,
        })

    context = "\n\n---\n\n".join(snippets) if snippets else "No context."

    # If Gemini LLM available, try LLM answer
    if key:
        llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            google_api_key=key,
            temperature=0.0,
        )

        prompt = (
            "Answer using ONLY the provided document context.\n"
            "Cite filename and page in your answer.\n\n"
            f"CONTEXT:\n{context}\n\nQUESTION:\n{question}\n\nANSWER:"
        )

        try:
            resp = llm.invoke(prompt)
            if hasattr(resp, "content"):
                return {"answer": resp.content, "sources": sources}
            return {"answer": str(resp), "sources": sources}
        except Exception as e:
            print("Gemini LLM failed; falling back:", e)
            traceback.print_exc()

    # Fallback answer (offline)
    fallback = (
        "Gemini not available; showing retrieved snippets instead:\n\n"
        + context
    )
    return {"answer": fallback, "sources": sources}
